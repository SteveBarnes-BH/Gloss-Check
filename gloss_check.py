#!/usr/bin/env python
#coding:utf-8
"""
  Purpose: Find the candidates for the glossary in a word document.
  Author: Steve Barnes --<StevenJohn.Barnes@ge.com>
  Created: 13/03/2017.
  Updated: 07/04/2017.

  Find the candidates for the glossary in a word document.
"""
#(Inspired by python-docx <https://github.com/mikemaccana/python-docx> &
#<http://etienned.github.io/posts/extract-text-from-word-docx-simply/>)
from __future__ import print_function
import sys
import os
import glob
import zipfile
import textwrap
import argparse
import codecs
import time
from collections import namedtuple
import tempfile

try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML

try:
    import wx
    GUI_OK = True
    __doc__ += '\n\nCall without argumets to start GUI.'
except ImportError:
    # wxPython not installed so no GUI available
    GUI_OK = False
    print('GUI not available without wxPython installation')

# Needs enchant to be installed for spell check.
try:
    import enchant
except ImportError:
    print('WARNING: No spell checker, install using: pip install pyenchant')
    enchant = None

# Docx will allow tables to be parsed.
try:
    import docx
except ImportError:
    print('WARNING: No docx so tables may be ommitted, install using: pip install docx')
    docx = None

try:
    import XXtextract
    USE_TEXTRACT = True
except ImportError:
    USE_TEXTRACT = False

from doc2docx import doc2docx
import text_utls

__version__ = "0.6-alpha"

USAGE = """
  Usage: gloss_check [-UCK] [-g <glossary_txt_file>] INPUT [...]
  Options:

    -U, --upper_only  Only consider all uppercase strings as candidates.
    -C, --chars_only: Exclude words with embedded numbers or symbols.
    -g <glossary_txt_file>, --existing_gloss=<glossary_txt_file>: An existing glossary to ignore
    -m, --inc_camel: Include any word with upper case after the first.
    -L, --lang=code Language code to spell check against [default: en_GB].
  """

# Constants used to decode MS Word Open Document Format
WRD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WRD_NS + 'p'
TEXT = WRD_NS + 't'
TABLE = WRD_NS + 'tbl'

# See if we have a dictionary available
if enchant is not None:
    LANGS = enchant.list_languages()
else:
    LANGS = []
# Command Line Arguments
ARG_LIST = [
    (['-M', '-m', '--min-acc'],
     {'type':int, "action":'store', "default":2, 'choices':range(2, 9),
      'help':"Minimum Length for a possible glossary entry"}),
    (['-u', '-U', '--upper_only'],
     {'action':'store_true',
      'help':"Only consider all uppercase strings as candidates."}),
    (['-C', '-c', '--chars_only'],
     {'action':'store_true',
      "help":"Exclude words with embedded numbers or symbols."}),
    (['-K', '-k', '--inc_camel',],
     {'action':'store_true',
      "help":'Include any word with upper any case after the first letter.'}),
    (['-o', '-1', '--one-per-line'],
     {'action':'store_true', "dest":'oneper',
      "help":"Display in a single column"}),
    (['-gu', '-GU', '--glossary-unused'],
     {'action':'store_true',
      "help":"Show unused entries from the glossary."}),
]

if len(LANGS): # Add the language option if we have a dictionary available
    ARG_LIST.insert(0, (['-L', '-l', '--lang'],
                        {'action':'store', 'choices':LANGS, 'default':'en_GB',
                         "help":'Language code to spell check against'}),
                   )

def get_textract_wordlist(path, minacc=1):
    """
    Take the path of a file as an argument, return the list of words.
    """
    try:
        wordlist = textract.process(path, 'utf-8').split()
        result = (True, text_utls.clean_wordlist(wordlist, minacc))
    except (WindowsError, textract.exceptions.ExtensionNotSupported):
        result = (False, [])
    return result

def get_doc_wordlist(path, minacc=1):  #, extract_glossary=False):
    """
    Take the path of a doc file as argument, return the list of words.
    """
    print("Need to convert .doc to temporary .docx")
    td = tempfile.gettempdir()
    temppath = os.path.join(td, 'gloss_check_temp.docx')
    if os.path.exists(temppath):
        os.remove(temppath)
    doc2docx(path, temppath)
    # Create a temprory copy of the doc as docx
    if docx is not None:
        ok, cwordlist = get_docx2_wordlist(temppath, minacc)
    else:
        ok, cwordlist = get_docx_wordlist(temppath, minacc)
    if os.path.exists(temppath):
        os.remove(temppath)
        print('Removed', temppath)
    return ok, cwordlist

def get_docx2_wordlist(path, minacc=1):
    """
    Use docx to get the word list.
    """
    if not os.path.splitext(path)[-1].lower() == '.docx':
        return False, []
    document = docx.Document(path)
    assert isinstance(document, docx.document.Document)
    wordlist = set()
    texts = [p.text for p in document.paragraphs]
    # Add the text from tables
    for table in document.tables:
        for col in table.columns:
            texts.extend([cell.text for cell in col.cells])
    if texts:
        text = '\n'.join(texts)
        wordlist.update(text.split())
    cwordlist = text_utls.clean_wordlist(wordlist, minacc)
    for word in get_docx_table_text(document, minacc=minacc):
        if word not in cwordlist:
            cwordlist.append(word)
    #document.lose()

    return len(cwordlist) > 0, sorted(cwordlist)

def get_docx_table_text(path_or_docx, minacc=1, tabno=None, colno=None):
    """
    Get the text from tables in a document supplied as a path or docx.
    If tabno is None then all tables, if it is a simple number then that one and
    if it is a list those.
    The same for the columns specified by colno.
    """
    texts = []
    if isinstance(path_or_docx, docx.document.Document):
        doc = path_or_docx
        #close_after = False
    else:
        doc = docx.Document(path_or_docx)
        #close_after = True
    if tabno is None:
        tables = doc.tables
    elif isinstance(tabno, int):
        tables = [doc.tables[tabno]]
    elif isinstance(tabno, list):
        tables = [doc.tables[no] for no in tabno]
    else:
        raise ValueError('get_docx_table_text tabno must be one of None, int or list of ints')
    if isinstance(colno, int):
        cols = [colno]
    elif isinstance(colno, list):
        cols = colno
    else:
        cols = None
    for table in tables:
        if colno is None:
            cols = range(len(table.columns))
        for n in cols:
            texts.extend([cell.text for cell in table.columns[n].cells])
    #if close_after:
        #doc.close()
    wordlist = set()
    if texts:
        text = '\n'.join(texts)
        wordlist.update(text.split())
    cwordlist = text_utls.clean_wordlist(wordlist, minacc)
    return cwordlist

def get_tree_table_wordlist(tree, minacc=1, tabno=None):
    """
    Get the wordlist from an XML tree.
    If tabno is None then all tables, if it is a simple number then that one and
    if it is a list those.
    """
    tables = []
    indexes = []
    texts = []
    for n, tab in enumerate(tree.getiterator(TABLE)):
        tables.append([node.text for node in tab.getiterator(TEXT)])
        indexes.append(n)
    if tabno is None:
        req_tabs = indexes
    elif isinstance(tabno, int):
        req_tabs = [tabno]
    elif isinstance(tabno, list) and all([isinstance(mem, int) for mem in tabno]):
        req_tabs = tabno
    else:
        raise ValueError("tabno must be one of None, int, list of ints")
    for n in req_tabs:
        texts.extend(tables[n])
    wordlist = set()
    if texts:
        text = ''.join(texts)
        wordlist.update(text.split())
    cwordlist = text_utls.clean_wordlist(wordlist, minacc)
    return cwordlist

def get_docx_wordlist(path, minacc=1):  #, extract_glossary=False):
    """
    Take the path of a docx file as argument, return the list of words.
    Note that this doesn't work as well as get_docx2_wordlist
    """
    if not os.path.splitext(path)[-1].lower() == '.docx':
        return False, []
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    wordlist = set()
    for paragraph in tree.getiterator(PARA):
        texts = [
            node.text
            for node in paragraph.getiterator(TEXT)
            if node.text]
        if texts:
            text = ''.join(texts)
            wordlist.update(text.split())
            if any([w in wordlist for w in suspect]):
                print('here')

    cwordlist = text_utls.clean_wordlist(wordlist, minacc)
    tabwords = get_tree_table_wordlist(tree, minacc=minacc)
    print(len(tabwords), 'words from tables.')
    for word in tabwords:
        if word not in cwordlist:
            cwordlist.append(word)

    return len(cwordlist) > 0, sorted(cwordlist)

def get_candidates(
        path, minlen=2, upper_only=True, chars_only=True, inc_cammel=False,
        existing_gloss=None, lang='en_GB'):
    """
    Get glossary candidates from a Word Open Document Format file.

    Params:
        path: The path to the docx file.
        upper_only: Only consider all upper case strings as candidates.
        inc_cammel: Include any word with upper case after the first.
        chars_only: Exclude words with embedded numbers or symbols.
        existing_gloss: An existing glossary to ignore.
        lang: Language code to spell check against.
    """
    success = False
    words = []
    method_dict = {'.docx': get_docx_wordlist, '.doc': get_doc_wordlist,}
    if docx is not None:
        method_dict['.docx'] = get_docx2_wordlist
    method = method_dict.get(os.path.splitext(path)[-1].lower())
    candiates = []
    unused = []
    if USE_TEXTRACT:
        success, words = get_textract_wordlist(path, minlen)
    elif method is not None:
        success, words = method(path, minlen)
    if success:
        candiates = get_candidates_from_list(
            words, upper_only, inc_cammel, chars_only, existing_gloss, lang)
        unused = [item for item in existing_gloss if not item in words]
    return (success, candiates, unused)

def get_candidates_from_list(words, upper_only=True, inc_cammel=False,
                             chars_only=True,
                             existing_gloss=None, lang='en_GB'):
    """
    Get glossary candidates from a word list.

    Params:
        words: The word list.
        upper_only: Only consider all upper case strings as candidates.
        inc_cammel: Include any word with upper case after the first.
        chars_only: Exclude words with embedded numbers or symbols.
        existing_gloss: An existing glossary to ignore.
        lang: Language code to spell check against.
    """
    if enchant is not None:
        try:
            chker = enchant.Dict(lang)
            words = [w for w in words if len(w) and not chker.check(w)]  # Not in dictionary
        except enchant.errors.DictNotFoundError:
            print(lang, "Dictionary Not Found, use -ll for list of options.")
            time.sleep(1.5)
    if chars_only:  # Only alpha based
        words = [w for w in words if all([c.isalpha() or c == '.' for c in w])]
    if upper_only:  # All upper only
        words = [w for w in words if all(
            [c.isupper() or c == '.' for c in w[:-1]]) and (w[-1].isupper() or w[-1] == 's')]
    if inc_cammel:  # All upper only
        words = [w for w in words if len(w) > 1 and any(
            [c.isupper() for c in w[1:]])]
    if existing_gloss:  # We have a glossary
        words = [w for w in words if w not in existing_gloss]

    return sorted(words)

def parse_args():
    """ Parse the arguments."""
    cl_args = [  # Items that only apply to the command line version.
        (['-G', '-g', '--glossary'],
         {'action':'append', "type":argparse.FileType('r'),
          "help":'An existing glossary to ignore'}),
        (['-v', '--version'],
         {'action':'store_true',
          "help":'Show version information and exit.'}),
        (['DOCS'],
         {"nargs":'*',
          'help':'Word document(s) to check, (Wildcards OK).'}),
    ]
    if len(LANGS):
        cl_args.append(
            (['-LL', '-ll', '--list-langs'],
             {'action':'store_true',
              "help":'List the available Language codes & exit'}))
    parser = argparse.ArgumentParser(
        #usage=USAGE,
        description=__doc__,
        #version='0.4'
    )
    for argp, argv in ARG_LIST:
        parser.add_argument(*argp, **argv)
    for argp, argv in cl_args:
        parser.add_argument(*argp, **argv)

    ops = parser.parse_args()
    if enchant is None:
        ops.lang = None
    if ops.version:
        show_versions()
    #print(ops)
    return ops

def show_versions():
    """ Show version information and exit."""
    print("Glossary Checker", __version__)
    if hasattr(sys, 'frozen'):
        print('Running Executable built from Python %s' % sys.version)
    else:
        print("Running under Python %s" % sys.version)
    if USE_TEXTRACT:
        print("Using: Textract!")
    if GUI_OK:
        print("Using: wx %s for GUI" % wx.version())
    else:
        print("No GUI! Try `pip install wxpython`")
    if enchant is not None:
        print("Using: pyEnchant V%s for spelling" % enchant.__version__)
    else:
        print("No Spell Checker! Try `pip install pyenchant`")
    if docx is not None:
        print("Using: python-docx %s for .DOCX parsing" % docx.__version__)
    #print("Using: Win32Com %s for .DOC conversions" % doc2docx.)

    sys.exit()

def get_glossary(ops):
    """ Get predefined glossaries."""
    ingloss = []
    if hasattr(ops, 'list_langs') and ops.list_langs:
        print(enchant.list_languages())
        sys.exit()
    if ops.glossary:
        print('Reading Glossary:')
        for item in ops.glossary:
            if isinstance(item, str) or sys.version_info[0] == 2 and isinstance(
                item, unicode):
                infile = open(item)
            else:
                infile = item
            print('   ', infile.name)
            ingloss.extend(infile.read().split())
            infile.close()
    ingloss = text_utls.clean_wordlist(ingloss)
    ingloss = get_candidates_from_list(
        ingloss, upper_only=ops.upper_only, chars_only=ops.chars_only,
        inc_cammel=ops.inc_camel, existing_gloss=[], lang=ops.lang)
    return ingloss

def process_docs(options, ext_gloss):
    """ Process the documents."""
    for arg in options.DOCS:
        print('Arg:', arg)
        for filename in glob.glob(arg):
            print('Processing', filename)
            success, candidates, unused = get_candidates(
                filename, minlen=options.min_acc, upper_only=options.upper_only,
                chars_only=options.chars_only, inc_cammel=options.inc_camel,
                existing_gloss=ext_gloss, lang=options.lang)
            if not success:
                print("ERROR: File is not a supported format or is corrupted/empty")
            else:
                print('%d Candidates:' % len(candidates))
                text_utls.smart_print(options, candidates)
                if options.glossary_unused:
                    print("Possible Unused Glossary Enties:")
                    text_utls.smart_print(options, unused)

def cmdline_main():
    """ Run the program."""
    ops = parse_args()
    ingloss = get_glossary(ops)
    process_docs(ops, ingloss)

if GUI_OK:
    class GuiPanel(wx.Panel):
        """ Main Window for the GUI."""
        def __init__(self, parent, log):
            super(GuiPanel, self).__init__(parent, -1)

            self.SetAutoLayout(True)
            outside_sizer = wx.BoxSizer(wx.VERTICAL)


            in_sizer = wx.BoxSizer(wx.HORIZONTAL)
            in_sizer.Add(FileDropPanel(self, log), 1, wx.EXPAND)

            outside_sizer.Add(in_sizer, 1, wx.EXPAND)
            self.SetSizer(outside_sizer)

    def start_gui():
        """ Start in GUI mode."""
        app = wx.App(redirect=False)
        frame = wx.Frame(None, -1, "Glossary Checker")
        win = GuiPanel(frame, sys.stdout)
        frame.Show(True)
        app.MainLoop()

    ##
    class MyFileDropTarget(wx.FileDropTarget):
        """
        File drop target allows the 'dropping' of files on the window to be
        handled.
        """
        ##
        def __init__(self, window, log):
            """ Initialise. """
            super(MyFileDropTarget, self).__init__()
            self.window = window
            self.log = log
            self.options = {'glossary':[],}

        ##
        def OnDropFiles(self, dummy_x, dummy_y, filenames):
            """
            Action files being dropped
            """
            self.window.clear_text()
            self.window.Refresh()
            self.window.set_insertion_point_end()
            self.window.write_text("\n%d file(s) dropped\n" %
                                   (len(filenames)))
            t_ops = namedtuple('options', self.options.keys())
            options = t_ops(**self.options)
            glossary = get_glossary(options)
            self.window.write_text("%d Glossary Entries Read!\n" % len(glossary))

            for name in filenames:
                self.window.write_text('\nProcessing %s!' % name)
                success, candidates, unused = get_candidates(
                    name, minlen=options.min_acc, upper_only=options.upper_only,
                    chars_only=options.chars_only, inc_cammel=options.inc_camel,
                    existing_gloss=glossary, lang=options.lang)
                if not success:
                    self.window.write_text(
                        " - ERROR: File is not supported or is corrupted/empty")
                else:
                    self.Smart_Write(
                        options, " %d Candidate Entries:\n" % len(candidates), candidates)
                    if options.glossary_unused:
                        self.Smart_Write(
                            options, '%d Unused Glossary Items:\n' % len(unused), unused)
            return 0

        def Smart_Write(self, options, title, items):
            """ Write the items based on the options."""
            self.window.write_text(title)
            if options.oneper:
                outtext = u'\n'.join(items)
            else:
                outtext = u'\n'.join(
                    textwrap.wrap(', '.join(items), 78))
            self.window.write_text(outtext + '\n')

    class FileDropPanel(wx.Panel):
        """ Class providing the actual file drop panel. """
        def __init__(self, parent, log):
            """ Initailse the Panel."""
            super(FileDropPanel, self).__init__(parent, -1)

            self.droptgt = MyFileDropTarget(self, log)
            sizer = wx.BoxSizer(wx.VERTICAL)
            sizer.Add(self.popultate_options(),
                      0, wx.EXPAND|wx.ALL, 1
                     )

            self.text = wx.TextCtrl(
                self, -1, "", style=wx.TE_MULTILINE|wx.HSCROLL|wx.TE_READONLY
            )

            self.text.SetDropTarget(self.droptgt)
            sizer.Add(self.text, 1, wx.EXPAND)
            self.text.WriteText("\nDrag one or more DOC/DOCX files here:")
            self.SetSizer(sizer)
            sizer.Fit(self)
            sizer.SetSizeHints(self)
            self.SetAutoLayout(True)

        def popultate_options(self):
            """ Add the options as selected from the command line arguments."""
            subsizer = wx.FlexGridSizer(0, 5, 0, 2)
            cb_style = wx.ALIGN_LEFT|wx.ALIGN_CENTER_VERTICAL|wx.ALL
            txt_style = wx.ALIGN_RIGHT|wx.ALIGN_CENTER_VERTICAL|wx.ALL
            ch_style = 0
            ctrl = wx.Button(self, -1, 'Load Glossary')
            ctrl.SetToolTip(
                wx.ToolTip('Select an existing glossary from text file(s) to ignore.'))
            self.Bind(wx.EVT_BUTTON, self.on_load_gloss, ctrl)
            subsizer.Add(ctrl, 1, cb_style, 2)

            for optnames, details in ARG_LIST:
                assert isinstance(details, dict)
                name = [n[2:] for n in optnames if n.startswith('--')][0]
                name = name.replace('-', '_')
                lable = name.replace('_', ' ').strip()
                lable = lable.capitalize()
                storeas = details.get('dest', name)
                if details['action'] in ['store_true', 'store_false']:
                    start_val = (details['action'] == 'store_false')
                    ctrl = wx.CheckBox(self, -1, lable, name=name)
                    self.droptgt.options[name] = start_val
                    ctrl.SetValue(start_val)
                    subsizer.Add(ctrl, 1, cb_style, 2)
                    self.Bind(wx.EVT_CHECKBOX, self.on_c_b, ctrl)
                elif details['action'] == 'store':
                    txt = wx.StaticText(self, -1, lable)
                    subsizer.Add(txt, 1, txt_style, 5)
                    ctrl = wx.Choice(self, -1, style=ch_style, name=name,
                                     choices=[str(c) for c in details['choices']])
                    start_val = details.get('default')
                    ctrl.SetStringSelection(str(start_val))
                    ctrl.rettype = details.get('type', str)
                    subsizer.Add(ctrl, 1, cb_style, 2)
                    self.Bind(wx.EVT_CHOICE, self.on_choose, ctrl)
                else:
                    ctrl = None
                    print(name, details)
                if ctrl:
                    ctrl.storeas = storeas
                    self.droptgt.options[storeas] = start_val
                    ctrl.SetToolTip(wx.ToolTip(details['help']))
            return subsizer
        ##
        def on_load_gloss(self, dummy_evt):
            """ Action the Load Glossary Button."""
            print('Load Glossary')
            self.droptgt.options['glossary'] = []
            dlg = wx.FileDialog(
                self, "Select Glossary File(s)", defaultDir='.',
                wildcard='Text|*.txt',
                style=wx.FD_OPEN|wx.FD_CHANGE_DIR|wx.FD_FILE_MUST_EXIST|wx.FD_MULTIPLE
                )
            if dlg.ShowModal() == wx.ID_OK:
                self.droptgt.options['glossary'] = dlg.GetPaths()
            dlg.Destroy()
            print(self.droptgt.options)

        def on_c_b(self, evt):
            """ Action a check box. """
            val = evt.EventObject.IsChecked()
            self.droptgt.options[evt.EventObject.storeas] = val
            print(self.droptgt.options)

        def on_choose(self, evt):
            """ Action a check box """
            val = evt.EventObject.GetStringSelection()
            stortype = evt.EventObject.rettype
            self.droptgt.options[evt.EventObject.storeas] = stortype(val)
            print(self.droptgt.options)

        ##
        def write_text(self, text):
            """ Wrapper for text.WriteText """
            self.text.WriteText(text)

        ##
        def set_insertion_point_end(self):
            """ Wrapper for text.SetInsertionPointEnd """
            self.text.SetInsertionPointEnd()

        ##
        def clear_text(self):
            """ Clear the text window """
            self.text.Clear()


if __name__ == '__main__':
    if GUI_OK and len(sys.argv) < 2:
        start_gui()
    else:
        cmdline_main()
