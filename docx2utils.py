#!/usr/bin/env python
#encoding:utf-8
"""
  Author:  Steve Barnes --<StevenJohn.Barnes@ge.com>
  Purpose: Docx based document utilities.
  Created: 17/04/2017
"""
from __future__ import (print_function, )

import os
try:
    import docx
except ImportError:
    docx = None

import text_utls

def get_docx2_wordlist(path, options):
    """
    Use docx to get the word list.
    """
    if not os.path.splitext(path)[-1].lower() == '.docx':
        return False, []
    document = docx.Document(path)
    assert isinstance(document, docx.document.Document)
    ignore_col1 = []
    poss_entries = []
    wordlist = set()
    texts = [p.text for p in document.paragraphs]
    if texts:
        text = '\n'.join(texts)
        wordlist.update(text_utls.tokenize(text, options))
    cwordlist = text_utls.clean_wordlist(wordlist, options.min_acc)

    if options.table_gloss:
        poss_entries, ignore_col1 = docx_get_table_gloss(path, options=options)
    # Get the table text from non-glossary tables
    non_gloss = [x for x in range(len(document.tables)) if x not in ignore_col1]
    for word in get_docx_table_text(document, non_gloss, options=options):
        if word not in cwordlist:
            cwordlist.append(word)
    # Then from the other colums of tables
    for word in get_docx_table_text(document, ignore_col1, excl_col=1, options=options):
        if word not in cwordlist:
            cwordlist.append(word)

    return (len(cwordlist) > 0, sorted(cwordlist, key=lambda s: s.lower()),
            sorted(poss_entries, key=lambda s: s.lower()))

def docx_get_table_gloss(path_or_docx, whitelist=None, options=None):
    """ Locate and parse tables that look like glossaries."""
    poss_entries = []
    ignore_col1 = []
    print("Searching for tables that might be glossaries.")
    (tables, col_nums, ignore_cols) = docx_table_text_valid_args(
        path_or_docx, None, 0, None)
    for index, table in enumerate(tables):
        coltxt = []
        for cell in table.columns[0].cells:
            coltxt.extend(text_utls.tokenize(cell.text, options))
        words = list(set(coltxt))
        candidates = text_utls.get_candidates_from_list(
            words, extern_gloss=whitelist, options=options)
        #print('Table %d has %d possible items of %d' % (
            #index, len(candidates), len(coltxt)))
        #print(', '.join(candidates))
        if len(candidates) * 100.0 / len(coltxt) > 50:  # %age possible
            print('Table %d has %d possible items of %d' % (
                index, len(candidates), len(coltxt)))
            print('Adding to assumed glossary:', ', '.join(candidates))
            ignore_col1.append(index)
            poss_entries.extend(candidates)
    return poss_entries, ignore_col1

def docx_table_text_valid_args(path_or_docx, tabno=None, colno=None,
                               excl_col=None):
    """ Deal with possible arguments."""
    col_nums = None
    ignore = []
    if isinstance(path_or_docx, docx.document.Document):
        doc = path_or_docx
    else:
        doc = docx.Document(path_or_docx)
    if tabno is None:
        tables = doc.tables
    elif isinstance(tabno, int):
        tables = [doc.tables[tabno]]
    elif isinstance(tabno, list):
        tables = [doc.tables[no] for no in tabno]
    else:
        raise ValueError(
            'get_docx_table_text tabno must be one of None, int or list of ints')
    if isinstance(colno, int):
        col_nums = [colno]
    elif isinstance(colno, list) and all([isinstance(item, int) for item in colno]):
        col_nums = colno
    elif colno is not None:
        print('get_docx_table_text colno must be one of None, int or list of ints')
        print('Ignoring', type(colno), colno)
    if isinstance(excl_col, int):
        ignore = []
    elif isinstance(excl_col, list) and all([isinstance(item, int) for item in excl_col]):
        ignore = excl_col
    elif excl_col is not None:
        print('get_docx_table_text excl_col must be int or list of ints')
        print('Ignoring', type(excl_col), excl_col)

    return (tables, col_nums, ignore)

def get_docx_table_text(path_or_docx, tabno=None, colno=None,
                        excl_col=None, options=None):
    """
    Get the text from tables in a document supplied as a path or docx.
    If tabno is None then all tables, if it is a simple number then that one and
    if it is a list those.
    The same for the columns specified by colno.
    """
    texts = []
    (tables, col_nums, ignore_cols) = docx_table_text_valid_args(
        path_or_docx, tabno, colno, excl_col)
    for table in tables:
        if colno is None:
            col_nums = range(len(table.columns))
        for col_number in col_nums:
            if col_number not in ignore_cols:
                texts.extend([cell.text for cell in table.columns[col_number].cells])
    #if close_after:
        #doc.close()
    wordlist = set()
    if texts:
        text = '\n'.join(texts)
        wordlist.update(text_utls.tokenize(text, options))
    cwordlist = text_utls.clean_wordlist(wordlist, options.min_acc)
    return cwordlist


if __name__ == '__main__':
    pass
